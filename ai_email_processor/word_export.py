"""
Word Document Export Module
Convert AI responses with Markdown formatting to Word documents
"""

import os
import re
from datetime import datetime
from typing import Dict, Any, Optional, List
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import markdown
    from markdown.extensions import codehilite, tables, toc
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False


class MarkdownToWordConverter:
    """Convert Markdown text to formatted Word document"""
    
    def __init__(self):
        self.doc = None
        self.styles_created = False
    
    def _check_dependencies(self) -> bool:
        """Check if required dependencies are available"""
        if not DOCX_AVAILABLE:
            print("Error: python-docx not available")
            print("Install with: pip install python-docx")
            return False
        
        return True
    
    def _create_custom_styles(self):
        """Create custom styles for document formatting"""
        if self.styles_created or not self.doc:
            return
        
        styles = self.doc.styles
        
        # Code block style
        try:
            code_style = styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
            code_font = code_style.font
            code_font.name = 'Courier New'
            code_font.size = Pt(10)
            code_font.color.rgb = RGBColor(0x00, 0x00, 0x80)  # Dark blue
            
            code_paragraph = code_style.paragraph_format
            code_paragraph.left_indent = Inches(0.5)
            code_paragraph.space_before = Pt(6)
            code_paragraph.space_after = Pt(6)
        except:
            pass  # Style might already exist
        
        # Quote style
        try:
            quote_style = styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
            quote_font = quote_style.font
            quote_font.italic = True
            quote_font.color.rgb = RGBColor(0x60, 0x60, 0x60)  # Gray
            
            quote_paragraph = quote_style.paragraph_format
            quote_paragraph.left_indent = Inches(0.5)
            quote_paragraph.space_before = Pt(6)
            quote_paragraph.space_after = Pt(6)
        except:
            pass
        
        # Emphasis styles
        try:
            emphasis_style = styles.add_style('Emphasis', WD_STYLE_TYPE.CHARACTER)
            emphasis_style.font.italic = True
        except:
            pass
        
        try:
            strong_style = styles.add_style('Strong', WD_STYLE_TYPE.CHARACTER)
            strong_style.font.bold = True
        except:
            pass
        
        self.styles_created = True
    
    def _parse_markdown_simple(self, text: str) -> List[Dict[str, Any]]:
        """Simple Markdown parser for basic formatting"""
        lines = text.split('\n')
        parsed_elements = []
        current_element = {'type': 'paragraph', 'content': ''}
        in_code_block = False
        code_block_content = []
        
        for line in lines:
            line_stripped = line.strip()
            
            # Code blocks
            if line_stripped.startswith('```'):
                if in_code_block:
                    # End code block
                    parsed_elements.append({
                        'type': 'code_block',
                        'content': '\n'.join(code_block_content)
                    })
                    code_block_content = []
                    in_code_block = False
                else:
                    # Start code block
                    if current_element['content'].strip():
                        parsed_elements.append(current_element)
                    current_element = {'type': 'paragraph', 'content': ''}
                    in_code_block = True
                continue
            
            if in_code_block:
                code_block_content.append(line)
                continue
            
            # Headers
            if line_stripped.startswith('# '):
                if current_element['content'].strip():
                    parsed_elements.append(current_element)
                parsed_elements.append({
                    'type': 'heading1',
                    'content': line_stripped[2:]
                })
                current_element = {'type': 'paragraph', 'content': ''}
            elif line_stripped.startswith('## '):
                if current_element['content'].strip():
                    parsed_elements.append(current_element)
                parsed_elements.append({
                    'type': 'heading2',
                    'content': line_stripped[3:]
                })
                current_element = {'type': 'paragraph', 'content': ''}
            elif line_stripped.startswith('### '):
                if current_element['content'].strip():
                    parsed_elements.append(current_element)
                parsed_elements.append({
                    'type': 'heading3',
                    'content': line_stripped[4:]
                })
                current_element = {'type': 'paragraph', 'content': ''}
            
            # Lists
            elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
                if current_element['content'].strip():
                    parsed_elements.append(current_element)
                parsed_elements.append({
                    'type': 'bullet_list_item',
                    'content': line_stripped[2:]
                })
                current_element = {'type': 'paragraph', 'content': ''}
            elif re.match(r'^\d+\. ', line_stripped):
                if current_element['content'].strip():
                    parsed_elements.append(current_element)
                parsed_elements.append({
                    'type': 'number_list_item',
                    'content': re.sub(r'^\d+\. ', '', line_stripped)
                })
                current_element = {'type': 'paragraph', 'content': ''}
            
            # Quotes
            elif line_stripped.startswith('> '):
                if current_element['content'].strip():
                    parsed_elements.append(current_element)
                parsed_elements.append({
                    'type': 'quote',
                    'content': line_stripped[2:]
                })
                current_element = {'type': 'paragraph', 'content': ''}
            
            # Horizontal rule
            elif line_stripped in ['---', '___', '***']:
                if current_element['content'].strip():
                    parsed_elements.append(current_element)
                parsed_elements.append({
                    'type': 'horizontal_rule',
                    'content': ''
                })
                current_element = {'type': 'paragraph', 'content': ''}
            
            # Empty line
            elif not line_stripped:
                if current_element['content'].strip():
                    parsed_elements.append(current_element)
                    current_element = {'type': 'paragraph', 'content': ''}
            
            # Regular paragraph content
            else:
                if current_element['content']:
                    current_element['content'] += '\n' + line
                else:
                    current_element['content'] = line
        
        # Add final element if it has content
        if current_element['content'].strip():
            parsed_elements.append(current_element)
        
        return parsed_elements
    
    def _apply_inline_formatting(self, paragraph, text: str):
        """Apply inline formatting like bold, italic, code to paragraph"""
        # This is a simplified approach - for production use, consider using a proper markdown parser
        
        # Split text by formatting markers while preserving them
        parts = []
        current = ""
        i = 0
        
        while i < len(text):
            # Bold (**text**)
            if i < len(text) - 1 and text[i:i+2] == '**':
                if current:
                    parts.append(('normal', current))
                    current = ""
                
                # Find closing **
                j = text.find('**', i + 2)
                if j != -1:
                    parts.append(('bold', text[i+2:j]))
                    i = j + 2
                    continue
                else:
                    current += text[i]
                    i += 1
            
            # Italic (*text*)
            elif text[i] == '*' and (i == 0 or text[i-1] != '*') and (i == len(text)-1 or text[i+1] != '*'):
                if current:
                    parts.append(('normal', current))
                    current = ""
                
                # Find closing *
                j = text.find('*', i + 1)
                if j != -1:
                    parts.append(('italic', text[i+1:j]))
                    i = j + 1
                    continue
                else:
                    current += text[i]
                    i += 1
            
            # Inline code (`text`)
            elif text[i] == '`':
                if current:
                    parts.append(('normal', current))
                    current = ""
                
                # Find closing `
                j = text.find('`', i + 1)
                if j != -1:
                    parts.append(('code', text[i+1:j]))
                    i = j + 1
                    continue
                else:
                    current += text[i]
                    i += 1
            
            else:
                current += text[i]
                i += 1
        
        if current:
            parts.append(('normal', current))
        
        # Add runs to paragraph with appropriate formatting
        for format_type, content in parts:
            run = paragraph.add_run(content)
            
            if format_type == 'bold':
                run.bold = True
            elif format_type == 'italic':
                run.italic = True
            elif format_type == 'code':
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)
    
    def create_document_from_markdown(self, content: str, title: str = "AI Response", 
                                    metadata: Dict[str, Any] = None) -> bool:
        """
        Create Word document from Markdown content
        
        Args:
            content (str): Markdown formatted content
            title (str): Document title
            metadata (Dict): Additional metadata (sender, timestamp, etc.)
            
        Returns:
            bool: True if document created successfully
        """
        if not self._check_dependencies():
            return False
        
        try:
            # Create new document
            self.doc = Document()
            self._create_custom_styles()
            
            # Add title
            title_paragraph = self.doc.add_heading(title, 0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata if provided
            if metadata:
                meta_paragraph = self.doc.add_paragraph()
                meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                meta_info = []
                if metadata.get('timestamp'):
                    meta_info.append(f"Generated: {metadata['timestamp']}")
                if metadata.get('sender'):
                    meta_info.append(f"Requested by: {metadata['sender']}")
                if metadata.get('provider'):
                    meta_info.append(f"AI Provider: {metadata['provider']}")
                if metadata.get('model'):
                    meta_info.append(f"Model: {metadata['model']}")
                
                meta_run = meta_paragraph.add_run(" | ".join(meta_info))
                meta_run.italic = True
                meta_run.font.size = Pt(10)
                meta_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            
            # Add separator
            self.doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Parse and add content
            elements = self._parse_markdown_simple(content)
            
            for element in elements:
                element_type = element['type']
                element_content = element['content']
                
                if element_type == 'heading1':
                    self.doc.add_heading(element_content, 1)
                
                elif element_type == 'heading2':
                    self.doc.add_heading(element_content, 2)
                
                elif element_type == 'heading3':
                    self.doc.add_heading(element_content, 3)
                
                elif element_type == 'code_block':
                    code_para = self.doc.add_paragraph()
                    try:
                        code_para.style = 'Code Block'
                    except:
                        # Fallback formatting
                        code_run = code_para.add_run(element_content)
                        code_run.font.name = 'Courier New'
                        code_run.font.size = Pt(10)
                    else:
                        code_para.add_run(element_content)
                
                elif element_type == 'quote':
                    quote_para = self.doc.add_paragraph()
                    try:
                        quote_para.style = 'Quote'
                        quote_para.add_run(element_content)
                    except:
                        # Fallback formatting
                        quote_run = quote_para.add_run(element_content)
                        quote_run.italic = True
                        quote_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
                
                elif element_type == 'bullet_list_item':
                    list_para = self.doc.add_paragraph()
                    list_para.style = 'List Bullet'
                    self._apply_inline_formatting(list_para, element_content)
                
                elif element_type == 'number_list_item':
                    list_para = self.doc.add_paragraph()
                    list_para.style = 'List Number'
                    self._apply_inline_formatting(list_para, element_content)
                
                elif element_type == 'horizontal_rule':
                    hr_para = self.doc.add_paragraph("_" * 50)
                    hr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                elif element_type == 'paragraph' and element_content.strip():
                    para = self.doc.add_paragraph()
                    self._apply_inline_formatting(para, element_content)
            
            return True
            
        except Exception as e:
            print(f"Error creating Word document: {e}")
            return False
    
    def save_document(self, filepath: str) -> bool:
        """
        Save the document to file
        
        Args:
            filepath (str): Path to save the document
            
        Returns:
            bool: True if saved successfully
        """
        if not self.doc:
            print("No document to save")
            return False
        
        try:
            # Ensure directory exists
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            
            # Save document
            self.doc.save(filepath)
            print(f"Document saved to: {filepath}")
            return True
            
        except Exception as e:
            print(f"Error saving document: {e}")
            return False


class WordExportManager:
    """Manager for Word document exports"""
    
    def __init__(self, export_dir: str = "ai_responses"):
        self.export_dir = export_dir
        self.ensure_export_directory()
    
    def ensure_export_directory(self):
        """Ensure export directory exists"""
        try:
            os.makedirs(self.export_dir, exist_ok=True)
        except Exception as e:
            print(f"Warning: Could not create export directory: {e}")
    
    def generate_filename(self, sender: str = None, timestamp: datetime = None,
                         content_type: str = None) -> str:
        """
        Generate appropriate filename for Word document
        
        Args:
            sender (str): Email sender
            timestamp (datetime): Generation timestamp
            content_type (str): Detected content type
            
        Returns:
            str: Generated filename
        """
        if timestamp is None:
            timestamp = datetime.now()
        
        # Clean sender for filename
        if sender:
            sender_clean = re.sub(r'[<>@\.\s]+', '_', sender)
            sender_clean = re.sub(r'[^\w\-_]', '', sender_clean)[:20]
        else:
            sender_clean = "unknown"
        
        # Create filename
        date_str = timestamp.strftime("%Y%m%d_%H%M%S")
        content_prefix = f"{content_type}_" if content_type and content_type != "generic" else ""
        
        filename = f"{content_prefix}ai_response_{sender_clean}_{date_str}.docx"
        
        return filename
    
    def export_ai_response(self, content: str, metadata: Dict[str, Any]) -> Optional[str]:
        """
        Export AI response to Word document
        
        Args:
            content (str): AI response content (markdown formatted)
            metadata (Dict): Response metadata
            
        Returns:
            str: Path to created file or None if failed
        """
        try:
            # Generate filename
            timestamp = datetime.now()
            filename = self.generate_filename(
                sender=metadata.get('sender'),
                timestamp=timestamp,
                content_type=metadata.get('content_type')
            )
            
            filepath = os.path.join(self.export_dir, filename)
            
            # Create converter and document
            converter = MarkdownToWordConverter()
            
            # Prepare title
            content_type = metadata.get('content_type', 'generic')
            title_parts = ["AI Response"]
            if content_type != 'generic':
                title_parts.append(f"({content_type.title()} Specialized)")
            
            title = " ".join(title_parts)
            
            # Prepare metadata for document
            doc_metadata = {
                'timestamp': timestamp.strftime("%Y-%m-%d %H:%M:%S"),
                'sender': metadata.get('sender'),
                'provider': metadata.get('provider_used'),
                'model': metadata.get('model'),
                'content_type': content_type,
                'is_superuser': metadata.get('is_superuser', False)
            }
            
            # Add super user info to title if applicable
            if metadata.get('is_superuser'):
                title = f"SUPER USER {title}"
            
            # Create document
            if converter.create_document_from_markdown(content, title, doc_metadata):
                if converter.save_document(filepath):
                    return filepath
            
            return None
            
        except Exception as e:
            print(f"Error exporting to Word: {e}")
            return None
    
    def get_export_statistics(self) -> Dict[str, Any]:
        """Get statistics about exported documents"""
        try:
            if not os.path.exists(self.export_dir):
                return {'total_exports': 0, 'files': []}
            
            files = [f for f in os.listdir(self.export_dir) if f.endswith('.docx')]
            
            file_info = []
            for filename in files:
                filepath = os.path.join(self.export_dir, filename)
                stat = os.stat(filepath)
                file_info.append({
                    'filename': filename,
                    'size': stat.st_size,
                    'modified': datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S')
                })
            
            # Sort by modification time (newest first)
            file_info.sort(key=lambda x: x['modified'], reverse=True)
            
            return {
                'total_exports': len(files),
                'export_directory': self.export_dir,
                'files': file_info[:10]  # Show last 10 files
            }
            
        except Exception as e:
            print(f"Error getting export statistics: {e}")
            return {'total_exports': 0, 'files': [], 'error': str(e)}


# Utility functions for integration

def export_response_to_word(content: str, metadata: Dict[str, Any], 
                          export_dir: str = "ai_responses") -> Optional[str]:
    """
    Quick function to export AI response to Word document
    
    Args:
        content (str): AI response content
        metadata (Dict): Response metadata
        export_dir (str): Export directory
        
    Returns:
        str: Path to created file or None if failed
    """
    manager = WordExportManager(export_dir)
    return manager.export_ai_response(content, metadata)


def check_word_export_requirements() -> Dict[str, bool]:
    """Check if Word export requirements are met"""
    return {
        'python_docx': DOCX_AVAILABLE,
        'markdown': MARKDOWN_AVAILABLE,
        'can_export': DOCX_AVAILABLE  # Markdown is optional for basic export
    }