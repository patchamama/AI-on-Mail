"""
Document Parser Module

This module handles parsing and text extraction from various document formats
including PDF and Word documents.
"""

import os
import tempfile
from typing import List, Dict, Optional, Any
from email.header import decode_header

# Optional imports for document processing
try:
    import PyPDF2
    from PyPDF2 import PdfReader
    PDF_SUPPORT = True
    PDF_LIBRARY = "PyPDF2"
except ImportError:
    try:
        import fitz  # PyMuPDF
        PDF_SUPPORT = True
        PDF_LIBRARY = "PyMuPDF"
    except ImportError:
        PDF_SUPPORT = False
        PDF_LIBRARY = None

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False


class DocumentParser:
    """Class for processing document attachments"""
    
    def __init__(self, max_size: int = 10000):
        """
        Initialize document processor
        
        Args:
            max_size (int): Maximum text size to extract in characters
        """
        self.max_size = max_size
        self.supported_formats = self._get_supported_formats()
    
    def _get_supported_formats(self) -> List[str]:
        """Get list of supported formats"""
        formats = []
        if PDF_SUPPORT:
            formats.extend(['pdf'])
        if DOCX_SUPPORT:
            formats.extend(['docx'])
        return formats
    
    def get_supported_formats(self) -> List[str]:
        """Return supported document formats"""
        return self.supported_formats.copy()
    
    def is_supported_format(self, filename: str) -> bool:
        """Check if file format is supported"""
        extension = filename.lower().split('.')[-1] if '.' in filename else ''
        return extension in self.supported_formats
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file
        
        Args:
            file_path (str): Path to PDF file
            
        Returns:
            str: Extracted text from PDF
        """
        if not PDF_SUPPORT:
            print("Warning: PDF processing not available")
            return ""
        
        text_content = ""
        
        try:
            if PDF_LIBRARY == "PyMuPDF":
                import fitz
                doc = fitz.open(file_path)
                for page in doc:
                    text_content += page.get_text()
                doc.close()
                print(f"Text extracted with PyMuPDF")
                
            elif PDF_LIBRARY == "PyPDF2":
                with open(file_path, 'rb') as file:
                    pdf_reader = PdfReader(file)
                    for page in pdf_reader.pages:
                        text_content += page.extract_text()
                print(f"Text extracted with PyPDF2")
                
        except Exception as e:
            print(f"Error extracting PDF text: {e}")
            return ""
        
        # Limit text size
        if len(text_content) > self.max_size:
            text_content = text_content[:self.max_size] + "...\n[Text truncated]"
        
        return text_content.strip()
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from Word document
        
        Args:
            file_path (str): Path to Word document
            
        Returns:
            str: Extracted text from document
        """
        if not DOCX_SUPPORT:
            print("Warning: Word document processing not available")
            return ""
        
        text_content = ""
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                text_content += "\n"
            
            print(f"Text extracted from Word document")
            
        except Exception as e:
            print(f"Error extracting Word text: {e}")
            return ""
        
        # Limit text size
        if len(text_content) > self.max_size:
            text_content = text_content[:self.max_size] + "...\n[Text truncated]"
        
        return text_content.strip()
    
    def decode_filename(self, filename: str) -> str:
        """
        Decode filename with special characters
        
        Args:
            filename (str): Filename to decode
            
        Returns:
            str: Decoded filename
        """
        try:
            decoded = decode_header(filename)[0]
            if isinstance(decoded[0], bytes):
                return decoded[0].decode(decoded[1] or 'utf-8')
            else:
                return decoded[0]
        except:
            return filename
    
    def process_attachment(self, part, temp_dir: str) -> Optional[Dict[str, Any]]:
        """
        Process individual attachment
        
        Args:
            part: Email message part
            temp_dir (str): Temporary directory for files
            
        Returns:
            Dict: Information about processed attachment or None
        """
        filename = part.get_filename()
        if not filename:
            return None
        
        # Decode filename
        filename = self.decode_filename(filename)
        
        # Check if format is supported
        if not self.is_supported_format(filename):
            print(f"Unsupported format: {filename}")
            return None
        
        try:
            # Save temporary file
            file_path = os.path.join(temp_dir, filename)
            with open(file_path, 'wb') as f:
                f.write(part.get_payload(decode=True))
            
            print(f"Downloaded: {filename}")
            
            # Extract text by type
            file_extension = filename.lower().split('.')[-1]
            text_content = ""
            
            if file_extension == 'pdf':
                text_content = self.extract_text_from_pdf(file_path)
            elif file_extension == 'docx':
                text_content = self.extract_text_from_docx(file_path)
            
            # Clean up temp file
            try:
                os.remove(file_path)
            except:
                pass
            
            if text_content:
                return {
                    'filename': filename,
                    'type': file_extension,
                    'content': text_content,
                    'size': len(text_content)
                }
            else:
                print(f"Could not extract content from: {filename}")
                return None
                
        except Exception as e:
            print(f"Error processing attachment {filename}: {e}")
            return None
    
    def process_email_attachments(self, msg) -> List[Dict[str, Any]]:
        """
        Process all email attachments
        
        Args:
            msg: Email message object
            
        Returns:
            List[Dict]: List of processed attachments
        """
        if not msg.is_multipart():
            return []
        
        processed_attachments = []
        temp_dir = tempfile.mkdtemp(prefix='email_attachments_')
        
        try:
            for part in msg.walk():
                content_disposition = str(part.get("Content-Disposition", ""))
                
                if "attachment" in content_disposition:
                    attachment_info = self.process_attachment(part, temp_dir)
                    if attachment_info:
                        processed_attachments.append(attachment_info)
                        print(f"Processed: {attachment_info['filename']} "
                              f"({attachment_info['type'].upper()}, "
                              f"{attachment_info['size']} characters)")
        
        except Exception as e:
            print(f"Error processing attachments: {e}")
        
        finally:
            # Clean up temp directory
            try:
                os.rmdir(temp_dir)
            except:
                pass
        
        return processed_attachments
